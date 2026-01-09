import re
import os
import csv
import threading
from docx import Document
from deep_translator import GoogleTranslator
from langdetect import detect
import time

# ================= Helper functions =================

def _normalise_position(prefix: str, num: str) -> str:
    first = num.split('.')[0]
    if prefix.strip().lower().startswith('art'):
        return 'Art. ' + first
    return '§' + first

def _clean_after_identifier(text: str) -> str:
    return re.sub(r'^[\s\)\.(]+', '', text).lstrip()

def _detect_language(text: str) -> str:
    if not text.strip():
        return "Unknown"
    try:
        lang = detect(text)
        if lang.startswith("pl"):
            return "Polish"
        elif lang.startswith("en"):
            return "English"
        else:
            return lang
    except Exception:
        return "Unknown"

# === 🧹 Nowa funkcja: usuwa tekst w nawiasach kwadratowych ===
def _remove_square_brackets(text: str) -> str:
    return re.sub(r'\[.*?\]', '', text).strip()


# === 🧠 Tłumaczenie dużych tekstów z podziałem na logiczne fragmenty ===
def _split_smart(text, max_chars=4000):
    if len(text) <= max_chars:
        return [text]
    parts = []
    while len(text) > max_chars:
        split_idx = max(
            text[:max_chars].rfind('. '),
            text[:max_chars].rfind(';'),
            text[:max_chars].rfind('\n')
        )
        if split_idx == -1 or split_idx < max_chars * 0.5:
            split_idx = max_chars
        parts.append(text[:split_idx+1])
        text = text[split_idx+1:]
    if text.strip():
        parts.append(text)
    return parts


def _translate_large_text(text: str, retry=5) -> str:
    if not text.strip():
        return ""

    lang = _detect_language(text)
    chunks = _split_smart(text, max_chars=4000)
    translated_chunks = []

    for chunk in chunks:
        for attempt in range(retry):
            try:
                if lang == "Polish":
                    translated = GoogleTranslator(source='pl', target='en').translate(chunk)
                elif lang == "English":
                    translated = GoogleTranslator(source='en', target='pl').translate(chunk)
                else:
                    translated = chunk
                translated_chunks.append(translated)
                break
            except Exception:
                time.sleep(2)
        else:
            translated_chunks.append("[Translation failed]")

    return ''.join(translated_chunks)


# ================= Core extraction =================

def extract_sections_from_docx(docx_path, csv_path, detect_language_flag=False, progress_callback=None):
    doc = Document(docx_path)
    filename = os.path.splitext(os.path.basename(docx_path))[0]

    pat_par = re.compile(r'^(?P<prefix>§)\s*(?P<num>\d+(?:\.\d+)*)')
    pat_art = re.compile(r'^(?P<prefix>Art\.)\s*(?P<num>\d+(?:\.\d+)*)', re.I)

    current_pos = None
    current_base = None
    bucket_lines = []
    rows = []

    def flush():
        nonlocal rows, bucket_lines, current_pos
        if current_pos is not None:
            description = '\n'.join(line for line in bucket_lines if line.strip())
            primary_lang = _detect_language(description) if detect_language_flag else ""
            translated = _translate_large_text(description)
            row = [
                description,
                current_pos,
                filename,
            ]
            if detect_language_flag:
                row.append(primary_lang)
            row.append(translated)
            # Dodatek nr 1
            row.append(current_pos)  # Position II
            # Dodatek nr 2
            row.append(f"{current_pos} {filename}")  # Position + Summary
            rows.append(row)
        bucket_lines = []
        current_pos = None

    total_paragraphs = len(doc.paragraphs)
    for idx, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        if not text:
            continue

        # 🧹 usuń nawiasy kwadratowe i ich zawartość
        text = _remove_square_brackets(text)

        m = pat_par.match(text) or pat_art.match(text)
        if m:
            prefix = m.group('prefix')
            num = m.group('num')
            top_num = num.split('.')[0]
            kind = '§' if prefix.startswith('§') else 'art'
            if current_base is None or (kind, top_num) != current_base:
                flush()
                current_base = (kind, top_num)
                current_pos = _normalise_position(prefix, num)
            rest = text[m.end():]
            rest = _clean_after_identifier(rest)
            if rest:
                bucket_lines.append(rest)
        else:
            if current_pos is not None:
                bucket_lines.append(text)
        if progress_callback:
            progress_callback(int((idx + 1) / total_paragraphs * 100))

    flush()

    with open(csv_path, mode='w', encoding='utf-8', newline='') as csvfile:
        writer = csv.writer(csvfile)
        header = ['Description', 'Position', 'Summary']
        if detect_language_flag:
            header.append('Primary Language')
        header.append('Description Translated')
        header.append('Position II')
        header.append('Position + Summary')
        writer.writerow(header)
        writer.writerows(rows)


# ================= GUI =================

def run_gui():
    import tkinter as tk
    from tkinter import filedialog, ttk, messagebox

    def start_extraction_thread():
        docx_path = filedialog.askopenfilename(
            title='Select a Word file',
            filetypes=[('Word files', '*.docx')]
        )
        if not docx_path:
            return
        csv_path = filedialog.asksaveasfilename(
            title='Save CSV as',
            defaultextension='.csv',
            filetypes=[('CSV files', '*.csv')]
        )
        if not csv_path:
            return

        progress_var.set(0)
        progress_bar.update()

        def progress_callback(percent):
            progress_var.set(percent)
            progress_bar.update()

        def run_extraction():
            try:
                extract_sections_from_docx(
                    docx_path,
                    csv_path,
                    detect_language_flag=var_langdetect.get(),
                    progress_callback=progress_callback
                )
                messagebox.showinfo('Done', f'Extracted sections into {csv_path}')
            except Exception as e:
                messagebox.showerror('Error', f'Extraction failed: {e}')

        threading.Thread(target=run_extraction, daemon=True).start()

    root = tk.Tk()
    root.title('Word → CSV Extractor for Jira')
    frame = ttk.Frame(root, padding=20)
    frame.grid(row=0, column=0)

    ttk.Label(frame, text='Convert Word file sections into Jira CSV').grid(row=0, column=0, columnspan=2, pady=10)

    var_translate = tk.BooleanVar(value=True)
    ttk.Checkbutton(frame, text='Auto Translate PL↔EN (always)', variable=var_translate, state='disabled').grid(row=1, column=0, sticky='w')

    var_langdetect = tk.BooleanVar(value=False)
    ttk.Checkbutton(frame, text='Detect Primary Language', variable=var_langdetect).grid(row=2, column=0, sticky='w')

    progress_var = tk.IntVar()
    progress_bar = ttk.Progressbar(frame, variable=progress_var, maximum=100, length=300)
    progress_bar.grid(row=3, column=0, pady=10)

    ttk.Button(frame, text='Start', command=start_extraction_thread).grid(row=4, column=0, pady=15)

    root.mainloop()


if __name__ == '__main__':
    run_gui()
